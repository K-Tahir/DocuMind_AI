# core/parser.py
import os
import fitz  # PyMuPDF
import pandas as pd
from sqlalchemy import create_engine, inspect, text
try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.docstore.document import Document

def extract_pdf_text(file_path: str) -> dict:
    """Reads a PDF page-by-page and returns structured text and pages."""
    print(f"[INFO] Parsing PDF: {file_path}")
    doc = fitz.open(file_path)
    pages = []
    full_text = []
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text_content = page.get_text()
        if text_content.strip():
            pages.append({
                "page": page_num + 1,
                "text": text_content
            })
            full_text.append(f"--- Page {page_num + 1} ---\n{text_content}")
            
    doc.close()
    print(f"[SUCCESS] Extracted {len(pages)} pages from PDF.")
    return {
        "type": "unstructured",
        "format": "pdf",
        "text": "\n\n".join(full_text),
        "pages": pages
    }

def extract_docx_text(file_path: str) -> dict:
    """Reads a Word document and returns text content including paragraphs and tables."""
    print(f"[INFO] Parsing DOCX: {file_path}")
    import docx
    doc = docx.Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Also extract tables and format them as Markdown-like text
    for i, table in enumerate(doc.tables):
        full_text.append(f"\n--- Table {i+1} ---")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_data))
            
    print(f"[SUCCESS] Extracted text from DOCX.")
    return {
        "type": "unstructured",
        "format": "docx",
        "text": "\n".join(full_text),
        "pages": []
    }

def extract_txt_text(file_path: str) -> dict:
    """Reads a plain text file."""
    print(f"[INFO] Parsing TXT: {file_path}")
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    print(f"[SUCCESS] Extracted text from TXT.")
    return {
        "type": "unstructured",
        "format": "txt",
        "text": content,
        "pages": []
    }

def extract_tabular_data(file_path: str) -> dict:
    """Reads CSV or Excel sheets using Pandas."""
    print(f"[INFO] Parsing tabular sheet: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".csv":
        df = pd.read_csv(file_path)
    else:
        df = pd.read_excel(file_path)
        
    # Standardize column names (strip whitespace)
    df.columns = [str(col).strip() for col in df.columns]
    
    # Convert dataframe to records
    records = df.to_dict(orient="records")
    
    # Also create a text summary of the tabular data for LLM processing if needed
    text_summary = f"Columns: {', '.join(df.columns)}\n"
    text_summary += f"Total rows: {len(df)}\n"
    text_summary += "Data sample:\n" + df.head(5).to_string()
    
    print(f"[SUCCESS] Extracted {len(records)} rows from tabular sheet.")
    return {
        "type": "tabular",
        "format": ext[1:],
        "data": records,
        "columns": list(df.columns),
        "text": text_summary
    }

def extract_database_table(connection_url: str, table_name: str) -> dict:
    """
    Connects to MySQL or SQLite, inspects schema, and queries rows.
    """
    print(f"[INFO] Parsing database table '{table_name}' from: {connection_url}")
    engine = create_engine(connection_url)
    
    # Inspect schema
    inspector = inspect(engine)
    columns_info = inspector.get_columns(table_name)
    schema = {col["name"]: str(col["type"]) for col in columns_info}
    
    # Fetch rows
    with engine.connect() as conn:
        result = conn.execute(text(f"SELECT * FROM {table_name} LIMIT 100"))
        # Get column names
        keys = result.keys()
        records = [dict(zip(keys, row)) for row in result]
        
    # Construct schema text summary
    schema_summary = f"Table: {table_name}\nSchema:\n"
    for col_name, col_type in schema.items():
        schema_summary += f" - {col_name}: {col_type}\n"
        
    print(f"[SUCCESS] Extracted schema and {len(records)} rows from table '{table_name}'.")
    return {
        "type": "relational",
        "format": "database",
        "table_name": table_name,
        "schema": schema,
        "data": records,
        "text": schema_summary + f"\nData sample (up to 100 rows):\n{records[:5]}"
    }

def parse_file(file_path: str, db_table_name: str = None, connection_url: str = None) -> dict:
    """Unified entrypoint for file parsing based on path or type."""
    if connection_url and db_table_name:
        return extract_database_table(connection_url, db_table_name)
        
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    elif ext == ".docx":
        return extract_docx_text(file_path)
    elif ext == ".txt":
        return extract_txt_text(file_path)
    elif ext in [".csv", ".xlsx", ".xls"]:
        return extract_tabular_data(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_pdf_tables(file_path: str) -> list:
    """Detects and extracts tables from PDF pages using PyMuPDF."""
    print(f"[INFO] Extracting tables from PDF: {file_path}")
    doc = fitz.open(file_path)
    tables = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        tabs = page.find_tables()
        for idx, t in enumerate(tabs):
            table_data = t.extract()
            if table_data:
                tables.append({
                    "page": page_num + 1,
                    "table_index": idx + 1,
                    "data": table_data
                })
    doc.close()
    print(f"[SUCCESS] Extracted {len(tables)} tables from PDF.")
    return tables

def extract_pdf_images(file_path: str, output_dir: str) -> list:
    """Extracts raw images/drawings from PDF pages and saves them to disk."""
    print(f"[INFO] Extracting figures from PDF: {file_path} to {output_dir}")
    doc = fitz.open(file_path)
    images_metadata = []
    os.makedirs(output_dir, exist_ok=True)
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                filename = f"figure_page{page_num+1}_{img_idx+1}.{image_ext}"
                filepath = os.path.join(output_dir, filename)
                with open(filepath, "wb") as f:
                    f.write(image_bytes)
                images_metadata.append({
                    "page": page_num + 1,
                    "image_index": img_idx + 1,
                    "filename": filename,
                    "path": filepath
                })
            except Exception as e:
                print(f"[WARNING] Failed to extract image at xref {xref}: {e}")
                
    doc.close()
    print(f"[SUCCESS] Extracted {len(images_metadata)} figures from PDF.")
    return images_metadata

def extract_docx_tables(file_path: str) -> list:
    """Extracts all tables from Word DOCX."""
    print(f"[INFO] Extracting tables from DOCX: {file_path}")
    import docx
    doc = docx.Document(file_path)
    tables = []
    for idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            "page": "N/A",
            "table_index": idx + 1,
            "data": table_data
        })
    print(f"[SUCCESS] Extracted {len(tables)} tables from DOCX.")
    return tables

def extract_docx_images(file_path: str, output_dir: str) -> list:
    """Extracts all embedded media images from Word DOCX and saves them to disk."""
    print(f"[INFO] Extracting figures from DOCX: {file_path} to {output_dir}")
    import docx
    doc = docx.Document(file_path)
    images_metadata = []
    os.makedirs(output_dir, exist_ok=True)
    img_idx = 1
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                img_bytes = rel.target_part.blob
                # Ext name from reference path
                img_ext = os.path.splitext(rel.target_ref)[1].replace(".", "")
                if not img_ext:
                    img_ext = "png"
                filename = f"figure_{img_idx}.{img_ext}"
                filepath = os.path.join(output_dir, filename)
                with open(filepath, "wb") as f:
                    f.write(img_bytes)
                images_metadata.append({
                    "page": "N/A",
                    "image_index": img_idx,
                    "filename": filename,
                    "path": filepath
                })
                img_idx += 1
            except Exception as e:
                print(f"[WARNING] Failed to extract image from DOCX rel: {e}")
                
    print(f"[SUCCESS] Extracted {len(images_metadata)} figures from DOCX.")
    return images_metadata